# auto_bot.py

import os
import time
import random
import logging
import re
import threading
import requests
from pathlib import Path
from dotenv import load_dotenv
from canvasapi import Canvas
from canvasapi.exceptions import CanvasException
from openai import OpenAI
from docx import Document
from concurrent.futures import ThreadPoolExecutor, as_completed
from bs4 import BeautifulSoup
import pypdf

from rich.console import Console
from rich.panel import Panel
from rich.progress import Progress, SpinnerColumn, BarColumn, TextColumn, TimeElapsedColumn
from rich.prompt import Confirm

load_dotenv()

# ── Console ───────────────────────────────────────────────────────────────────
console = Console()

# ── Configuration ─────────────────────────────────────────────────────────────
BASE_DIR     = Path(__file__).parent
OUTPUT_DIR   = BASE_DIR / "Completed_Homework"
CANVAS_URL   = os.getenv("CANVAS_URL", "").rstrip("/")
CANVAS_TOKEN = os.getenv("CANVAS_TOKEN", "")
AI_MODEL     = os.getenv("AI_MODEL", "mistral-nemo")
STUDENT_NAME = os.getenv("STUDENT_NAME", "a student")
OLLAMA_URL   = os.getenv("OLLAMA_URL", "http://localhost:11434")

IGNORE_LIST   = ["Test", "Quiz", "Final Exam", "Physical Education", "Spartan Central", "Industrial Tech"]
PRIORITY_LIST = []

# Max concurrent assignment downloads per course (keep low for stealth)
ASSIGNMENT_WORKERS = 2

# Realistic browser User-Agent used on every outbound HTTP request
BROWSER_UA = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/124.0.0.0 Safari/537.36"
)

# Stealth delay ranges (seconds). All intervals are random uniform draws.
_DELAYS = {
    "between_courses":     (10.0, 25.0),  # pause between finishing one course and starting the next
    "between_assignments": (2.0,  5.0),   # pause before dispatching each assignment to the thread pool
    "after_bulk_fetch":    (1.5,  3.5),   # pause after calling get_assignments() on a course
    "after_api_call":      (0.5,  1.5),   # pause after individual Canvas API calls (submission check)
    "before_download":     (1.0,  2.5),   # pause before each HTTP file download
}

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

client = OpenAI(base_url=f"{OLLAMA_URL}/v1", api_key="ollama")

# Serialises confirmation prompts so only one appears at a time across threads
_confirm_lock = threading.Lock()

logging.basicConfig(
    filename=str(BASE_DIR / "sys_check.log"),
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
)

# ── Stealth utilities ─────────────────────────────────────────────────────────

def _sleep(key: str) -> None:
    lo, hi = _DELAYS[key]
    time.sleep(random.uniform(lo, hi))

def call_with_backoff(fn, *args, max_retries: int = 3, base_delay: float = 5.0, **kwargs):
    """
    Call fn(*args, **kwargs) with exponential backoff on Canvas rate-limit errors.
    Returns fn's result, or None if all retries are exhausted.
    """
    for attempt in range(max_retries + 1):
        try:
            return fn(*args, **kwargs)
        except (CanvasException, requests.exceptions.HTTPError) as exc:
            msg = str(exc).lower()
            status = getattr(getattr(exc, "response", None), "status_code", None)
            is_rate_limit = "429" in msg or "too many requests" in msg or status == 429
            if is_rate_limit and attempt < max_retries:
                wait = base_delay * (2 ** attempt) + random.uniform(0, 2)
                console.log(f"[yellow]⚠ Rate limited — waiting {wait:.0f}s (attempt {attempt + 1}/{max_retries})[/yellow]")
                logging.warning(f"Rate limited, backing off {wait:.0f}s")
                time.sleep(wait)
            else:
                raise
    return None

# ── Core utilities ────────────────────────────────────────────────────────────

def sanitize_filename(name: str) -> str:
    return re.sub(r"[^\w\-_.() ]", "", name).strip()

def is_valid_course(course) -> bool:
    name = getattr(course, "name", "Unknown")
    if any(word.lower() in name.lower() for word in IGNORE_LIST):
        return False
    if getattr(course, "access_restricted_by_date", False):
        return False
    return True

def extract_links_from_description(html_desc: str) -> list[str]:
    """Parse all downloadable file/document links from an assignment's HTML description."""
    if not html_desc:
        return []
    soup = BeautifulSoup(html_desc, "html.parser")
    links = []
    for a in soup.find_all("a", href=True):
        href = a["href"]
        if any(x in href for x in [
            "docs.google.com/document",
            "/files/",
            "/courses/",
            ".pdf",
            ".docx",
            ".doc",
        ]):
            links.append(href)
    return links

def get_assignment_attachments(assignment) -> list[str]:
    """Return download URLs for files attached directly to the assignment."""
    urls = []
    for att in getattr(assignment, "attachments", []) or []:
        url = att.get("url") or att.get("preview_url")
        if url:
            urls.append(url)
    return urls

# ── Download layer ────────────────────────────────────────────────────────────

def download_gdoc_as_docx(url: str, safe_filename: str):
    """Export a Google Doc to DOCX and save it as a temp file."""
    _sleep("before_download")
    try:
        doc_id_match = re.search(r"/d/([a-zA-Z0-9_-]+)", url)
        if not doc_id_match:
            logging.warning(f"Invalid Google Doc URL: {url}")
            return None
        export_url = f"https://docs.google.com/document/d/{doc_id_match.group(1)}/export?format=docx"
        resp = requests.get(export_url, headers={"User-Agent": BROWSER_UA}, timeout=20)
        if resp.status_code == 200:
            tmp = OUTPUT_DIR / f"TEMP_{random.randint(1000, 9999)}_{safe_filename}.docx"
            tmp.write_bytes(resp.content)
            return tmp
    except Exception as exc:
        logging.error(f"Google Doc download error for {url}: {exc}")
    return None

def download_canvas_file(url: str, safe_filename: str):
    """Download a Canvas-hosted file using the API token for auth."""
    _sleep("before_download")
    try:
        headers = {
            "Authorization": f"Bearer {CANVAS_TOKEN}",
            "User-Agent": BROWSER_UA,
        }
        resp = requests.get(url, headers=headers, timeout=30, allow_redirects=True)
        if resp.status_code != 200:
            logging.warning(f"Canvas file download returned {resp.status_code} for {url}")
            return None
        ct = resp.headers.get("content-type", "")
        if "pdf" in ct or url.lower().endswith(".pdf"):
            ext = ".pdf"
        elif "word" in ct or url.lower().endswith((".docx", ".doc")):
            ext = ".docx"
        else:
            ext = ".bin"
        tmp = OUTPUT_DIR / f"TEMP_{random.randint(1000, 9999)}_{safe_filename}{ext}"
        tmp.write_bytes(resp.content)
        return tmp
    except Exception as exc:
        logging.error(f"Canvas file download error for {url}: {exc}")
    return None

# ── Text extraction ───────────────────────────────────────────────────────────

def extract_text_from_docx(path) -> str | None:
    try:
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip()) or None
    except Exception as exc:
        logging.error(f"DOCX extraction error for {path}: {exc}")
        return None

def extract_text_from_pdf(path) -> str | None:
    try:
        reader = pypdf.PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages).strip() or None
    except Exception as exc:
        logging.error(f"PDF extraction error for {path}: {exc}")
        return None

def extract_text(path) -> str | None:
    suffix = Path(path).suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    if suffix in (".docx", ".doc"):
        return extract_text_from_docx(path)
    return None

# ── AI persona engine ─────────────────────────────────────────────────────────

def ai_fill_worksheet(worksheet_text: str) -> str | None:
    prompt = f"You are {STUDENT_NAME}, a 15-year-old 9th-grade student.\n"
    prompt += "- Use natural, slightly casual vocabulary. Use phrases like 'I think' or 'Basically'.\n"
    prompt += "- Avoid overly formal AI words like 'moreover' or 'delve'.\n"
    prompt += "- Provide perfect, highly detailed, and completely accurate answers.\n"
    prompt += "- ONLY return the answers in the exact same format as the original document.\n"
    prompt += "- Do NOT add any extra formatting, asterisks, or notes.\n"
    prompt += "- Answer each question separately in the same order.\n"
    prompt += "- Put answers directly below each question in the same font, size, and formatting.\n"
    prompt += "- Keep all question text exactly as it appears in the original.\n"
    prompt += "- Only fill in the blank spaces or answer areas where appropriate.\n\n"
    prompt += f"Worksheet Text:\n{worksheet_text}"

    try:
        response = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model=AI_MODEL,
            temperature=0.3,
        )
        return response.choices[0].message.content
    except Exception as exc:
        logging.error(f"AI service error: {exc}")
        return None

# ── Canvas submission ─────────────────────────────────────────────────────────

def is_already_submitted(assignment) -> bool:
    try:
        sub = call_with_backoff(assignment.get_submission, "self")
        return getattr(sub, "workflow_state", "") in ("submitted", "graded")
    except Exception:
        return False

def upload_file_for_submission(course_id, assignment_id, file_path):
    """
    Three-step Canvas file upload. Returns the file ID on success, else None.
    Steps 1→2→3 are intentionally kept together without delays (atomic sequence).
    """
    file_path = Path(file_path)
    ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    headers = {
        "Authorization": f"Bearer {CANVAS_TOKEN}",
        "User-Agent": BROWSER_UA,
    }

    # Step 1 — notify Canvas of the upcoming upload
    init_url = f"{CANVAS_URL}/api/v1/courses/{course_id}/assignments/{assignment_id}/submissions/self/files"
    init = requests.post(init_url, headers=headers, json={
        "name":         file_path.name,
        "size":         file_path.stat().st_size,
        "content_type": ct,
    })
    if init.status_code != 200:
        logging.error(f"Upload initiation failed ({init.status_code}): {init.text}")
        return None

    info = init.json()

    # Step 2 — PUT the bytes to the pre-signed URL
    with open(file_path, "rb") as f:
        up = requests.post(
            info["upload_url"],
            data=info.get("upload_params", {}),
            files={"file": (file_path.name, f, ct)},
            allow_redirects=False,
        )

    # Step 3 — confirm (Canvas may redirect to the final file record)
    if up.status_code in (301, 302):
        confirm = requests.get(up.headers["Location"], headers=headers)
        file_info = confirm.json()
    elif up.status_code in (200, 201):
        file_info = up.json()
    else:
        logging.error(f"File upload step 2 failed ({up.status_code})")
        return None

    return file_info.get("id")

def confirm_submission(assignment, course_name: str, solved_text: str, progress: Progress) -> bool:
    """
    Show a rich panel preview and ask the user whether to submit to Canvas.
    Serialised by _confirm_lock so parallel threads queue up one at a time.
    Temporarily pauses the live progress display while reading input.
    """
    types = getattr(assignment, "submission_types", []) or []
    preview = " ".join(solved_text.split())[:550]
    if len(preview) == 550:
        preview += "…"

    content = (
        f"[bold white]{assignment.name}[/bold white]\n"
        f"[dim]Course:[/dim]  {course_name}\n"
        f"[dim]Type:[/dim]    {', '.join(types) or 'unknown'}\n\n"
        f"[italic dim]{preview}[/italic dim]"
    )

    with _confirm_lock:
        progress.stop()
        console.print()
        console.print(Panel(
            content,
            title="[yellow bold]  Confirm Submission  [/yellow bold]",
            border_style="yellow",
            padding=(1, 3),
        ))
        result = Confirm.ask("  Submit this to Canvas?", default=False, console=console)
        console.print()
        progress.start()

    return result

def submit_assignment(assignment, course_id, solved_text: str, file_path=None) -> bool:
    """Submit completed work to Canvas based on the assignment's submission type."""
    types = getattr(assignment, "submission_types", []) or []

    try:
        if "online_text_entry" in types:
            html_body = "<p>" + solved_text.replace("\n", "</p><p>") + "</p>"
            call_with_backoff(assignment.submit, submission={
                "submission_type": "online_text_entry",
                "body": html_body,
            })
            logging.info(f"[SUBMITTED text] {assignment.name}")
            return True

        elif "online_upload" in types and file_path:
            file_id = upload_file_for_submission(course_id, assignment.id, file_path)
            if file_id:
                call_with_backoff(assignment.submit, submission={
                    "submission_type": "online_upload",
                    "file_ids": [file_id],
                })
                logging.info(f"[SUBMITTED file] {assignment.name}")
                return True
            logging.warning(f"[UPLOAD FAILED] {assignment.name} — saved locally only")

        else:
            logging.info(f"[LOCAL ONLY] types={types} — {assignment.name}")

    except Exception as exc:
        logging.error(f"Submission error for {assignment.name}: {exc}")

    return False

# ── Document builder ──────────────────────────────────────────────────────────

def build_output_docx(original_path, solved_text: str, out_path) -> None:
    """Mirror the original's paragraph styles but replace text with AI output."""
    original_doc = Document(original_path)
    solved_paras = solved_text.split("\n")
    orig_paras   = original_doc.paragraphs
    new_doc      = Document()

    for i, orig_para in enumerate(orig_paras):
        new_para = new_doc.add_paragraph()
        new_para.style = orig_para.style
        new_para.add_run(solved_paras[i] if i < len(solved_paras) else orig_para.text)

    for j in range(len(orig_paras), len(solved_paras)):
        if solved_paras[j].strip():
            new_doc.add_paragraph(solved_paras[j])

    new_doc.save(out_path)

def build_plain_docx(solved_text: str, out_path) -> None:
    doc = Document()
    for line in solved_text.split("\n"):
        doc.add_paragraph(line)
    doc.save(out_path)

# ── Assignment processor ──────────────────────────────────────────────────────

def process_assignment(course_folder, course_name: str, assignment, course_id,
                       progress: Progress, assign_task_id) -> None:
    safe_assign = sanitize_filename(assignment.name)

    # Check if already submitted (Canvas API call — add small delay after)
    already = is_already_submitted(assignment)
    _sleep("after_api_call")

    if already:
        console.log(f"[dim]↷ Already submitted:[/dim] {assignment.name}")
        logging.info(f"[SKIP submitted] {safe_assign}")
        progress.update(assign_task_id, advance=1)
        return

    desc  = getattr(assignment, "description", "") or ""
    links = extract_links_from_description(desc) + get_assignment_attachments(assignment)

    if not links:
        progress.update(assign_task_id, advance=1)
        return

    logging.info(f"[SOLVING] {assignment.name} — {len(links)} file(s)")

    for i, url in enumerate(links, start=1):
        label      = f"[DONE] {safe_assign}_part{i}"
        final_path = course_folder / f"{label}.docx"

        if final_path.exists():
            console.log(f"[dim]↷ Already done:[/dim] {assignment.name} (part {i})")
            continue

        temp_path = None
        try:
            # Download
            console.log(f"[cyan]↓ Downloading:[/cyan] {assignment.name} (part {i})")
            if "docs.google.com/document" in url:
                temp_path = download_gdoc_as_docx(url, label)
            else:
                temp_path = download_canvas_file(url, label)

            if not temp_path:
                console.log(f"[red]✗ Download failed:[/red] {url}")
                continue

            # Extract text
            text = extract_text(temp_path)
            if not text:
                console.log(f"[red]✗ No text extracted:[/red] {url}")
                continue

            # AI solve
            console.log(f"[yellow]⟳ Solving with AI:[/yellow] {assignment.name}")
            solved = ai_fill_worksheet(text)
            if not solved:
                continue

            # Save locally (always, regardless of submission choice)
            if temp_path.suffix == ".docx":
                build_output_docx(temp_path, solved, final_path)
            else:
                build_plain_docx(solved, final_path)

            console.log(f"[green]✓ Saved:[/green] {label}")
            logging.info(f"[SAVED] {course_name} -> {label}")

            # Human confirmation before Canvas submission
            if confirm_submission(assignment, course_name, solved, progress):
                _sleep("after_api_call")
                submitted = submit_assignment(assignment, course_id, solved, final_path)
                if submitted:
                    console.log(f"[green bold]↑ Submitted:[/green bold] {assignment.name}")
                    logging.info(f"[SUCCESS] {course_name} -> {label}")
                else:
                    console.log(f"[yellow]⚠ Submit failed — file kept locally[/yellow]")
            else:
                console.log(f"[dim]↷ Skipped submission:[/dim] {assignment.name}")
                logging.info(f"[SKIPPED by user] {label}")

        except Exception as exc:
            logging.error(f"Error processing {label}: {exc}")
            console.log(f"[red]✗ Error:[/red] {assignment.name} — {exc}")
        finally:
            if temp_path and Path(temp_path).exists():
                Path(temp_path).unlink()

    progress.update(assign_task_id, advance=1)

# ── Course processor ──────────────────────────────────────────────────────────

def process_course(course, progress: Progress, course_task_id, assign_task_id) -> None:
    course_name = getattr(course, "name", "Unknown Course")
    console.rule(f"[bold cyan]{course_name}[/bold cyan]")

    safe_course  = sanitize_filename(course_name)
    course_folder = OUTPUT_DIR / safe_course
    course_folder.mkdir(parents=True, exist_ok=True)

    try:
        assignments = call_with_backoff(lambda: list(course.get_assignments()))
    except Exception as exc:
        logging.error(f"Could not fetch assignments for {course_name}: {exc}")
        console.log(f"[red]✗ Failed to fetch assignments:[/red] {course_name}")
        progress.update(course_task_id, advance=1)
        return

    _sleep("after_bulk_fetch")

    # Filter to assignments that have something to download
    actionable = [
        a for a in (assignments or [])
        if not any(w.lower() in a.name.lower() for w in IGNORE_LIST)
        and (
            extract_links_from_description(getattr(a, "description", "") or "")
            or get_assignment_attachments(a)
        )
    ]

    console.log(f"[dim]{len(actionable)} actionable assignment(s) found[/dim]")

    # Grow the assignments bar total now that we know the count
    current_total = progress.tasks[assign_task_id].total or 0
    progress.update(assign_task_id, total=current_total + len(actionable))

    futures = []
    with ThreadPoolExecutor(max_workers=ASSIGNMENT_WORKERS) as executor:
        for assignment in actionable:
            _sleep("between_assignments")   # pace dispatching — key stealth measure
            futures.append(executor.submit(
                process_assignment,
                course_folder, course_name, assignment, course.id,
                progress, assign_task_id,
            ))

        for future in as_completed(futures):
            try:
                future.result()
            except Exception as exc:
                logging.error(f"Assignment thread error: {exc}")

    progress.update(course_task_id, advance=1)

# ── Main ──────────────────────────────────────────────────────────────────────

def main() -> None:
    if not CANVAS_URL or not CANVAS_TOKEN:
        console.print(Panel(
            "[bold red]Canvas URL or token not configured.[/bold red]\n"
            "Run [cyan]python run.py[/cyan] and choose Setup first.",
            border_style="red",
        ))
        return

    console.print(Panel(
        f"[bold cyan]AutoCanvas[/bold cyan]\n"
        f"[dim]Student:[/dim] [green]{STUDENT_NAME}[/green]  "
        f"[dim]Model:[/dim] [yellow]{AI_MODEL}[/yellow]  "
        f"[dim]Ollama:[/dim] {OLLAMA_URL}",
        border_style="cyan",
    ))

    try:
        canvas      = Canvas(CANVAS_URL, CANVAS_TOKEN)
        user        = canvas.get_current_user()
        raw_courses = call_with_backoff(lambda: list(user.get_courses(enrollment_state="active")))
        courses     = [c for c in (raw_courses or []) if is_valid_course(c)]
        courses.sort(key=lambda c: 0 if any(p.lower() in c.name.lower() for p in PRIORITY_LIST) else 1)
        console.log(f"Logged in as [bold]{getattr(user, 'name', 'Unknown')}[/bold] — {len(courses)} active course(s)")
    except Exception as exc:
        logging.critical(f"Initialization failed: {exc}")
        console.print(f"[bold red]CRITICAL ERROR:[/bold red] {exc}")
        return

    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}"),
        BarColumn(),
        TextColumn("[progress.completed]{task.completed}[dim]/[/dim]{task.total}"),
        TimeElapsedColumn(),
        console=console,
    ) as progress:
        course_task = progress.add_task("[cyan]Courses[/cyan]",     total=len(courses))
        assign_task = progress.add_task("[green]Assignments[/green]", total=0)

        for course in courses:
            process_course(course, progress, course_task, assign_task)
            _sleep("between_courses")   # most impactful stealth delay

    console.rule("[bold green]All Classes Complete[/bold green]")

if __name__ == "__main__":
    main()
